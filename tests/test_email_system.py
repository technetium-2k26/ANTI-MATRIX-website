import unittest
import os
import docx
from datetime import datetime, timezone
from app import create_app
from models import db, User, JobPosting, JobApplication, Employee, EmployeeDocument, DocumentTemplate, EmailTemplate, EmailLog
from services.email_service import (
    DEFAULT_APPLICATION_SUCCESSFUL_SUBJECT,
    DEFAULT_APPLICATION_SUCCESSFUL_BODY,
    DEFAULT_OFFER_LETTER_SUBJECT,
    DEFAULT_OFFER_LETTER_BODY,
    markdown_to_html_email,
    replace_variables,
    send_application_successful_email,
    send_offer_letter_shortlisted_email,
    render_sample_email_preview,
    send_test_email
)
from services.offer_letter_service import generate_offer_letter_docx


class TestAntiMatrixEmailSystem(unittest.TestCase):
    def setUp(self):
        """Set up test application with in-memory database."""
        self.app = create_app('testing')
        self.app.config['WTF_CSRF_ENABLED'] = False

        self.client = self.app.test_client()
        self.ctx = self.app.app_context()
        self.ctx.push()

        db.create_all()

        # Seed or fetch Admin User
        self.admin = User.query.filter_by(email='admin@antimatrix.ai').first()
        if not self.admin:
            self.admin = User(name='Test Admin', email='admin@antimatrix.ai', role='admin', is_active=True)
            self.admin.set_password('Admin@AntiMatrix2026!')
            db.session.add(self.admin)

        # Seed or fetch Member User
        self.member = User.query.filter_by(email='member@example.com').first()
        if not self.member:
            self.member = User(name='Regular Member', email='member@example.com', role='member', is_active=True)
            self.member.set_password('Member@2026!')
            db.session.add(self.member)

        # Seed DocumentTemplate for Offer Letter
        templates_dir = os.path.join(self.app.root_path, 'uploads', 'templates')
        os.makedirs(templates_dir, exist_ok=True)
        master_template_path = os.path.join(templates_dir, 'email_test_offer_letter_master.docx')
        if not os.path.exists(master_template_path):
            doc = docx.Document()
            doc.add_paragraph("Anti-Matrix Master Offer Letter")
            doc.add_paragraph("Dear [Candidate Name], Congratulations on [Job Title]. Ref: [Reference Number]")
            doc.save(master_template_path)

        doc_tmpl = DocumentTemplate.query.filter_by(template_type='offer_letter').first()
        if not doc_tmpl:
            doc_tmpl = DocumentTemplate(
                template_type='offer_letter',
                name='Anti-Matrix Master Offer Letter',
                filename='email_test_offer_letter_master.docx',
                file_path=master_template_path,
                is_active=True
            )
            db.session.add(doc_tmpl)

        # Seed or fetch Official Email Templates
        self.app_success_tmpl = EmailTemplate.query.filter_by(template_type='application_successful').first()
        if not self.app_success_tmpl:
            self.app_success_tmpl = EmailTemplate(
                template_type='application_successful',
                name='Application Successful',
                subject=DEFAULT_APPLICATION_SUCCESSFUL_SUBJECT,
                body=DEFAULT_APPLICATION_SUCCESSFUL_BODY
            )
            db.session.add(self.app_success_tmpl)
        else:
            self.app_success_tmpl.subject = DEFAULT_APPLICATION_SUCCESSFUL_SUBJECT
            self.app_success_tmpl.body = DEFAULT_APPLICATION_SUCCESSFUL_BODY

        self.offer_tmpl = EmailTemplate.query.filter_by(template_type='offer_letter').first()
        if not self.offer_tmpl:
            self.offer_tmpl = EmailTemplate(
                template_type='offer_letter',
                name='Offer Letter / Shortlisted',
                subject=DEFAULT_OFFER_LETTER_SUBJECT,
                body=DEFAULT_OFFER_LETTER_BODY
            )
            db.session.add(self.offer_tmpl)
        else:
            self.offer_tmpl.subject = DEFAULT_OFFER_LETTER_SUBJECT
            self.offer_tmpl.body = DEFAULT_OFFER_LETTER_BODY

        # Create Job Opening
        self.job = JobPosting(
            title='AI Engineer Intern',
            department='Engineering',
            location='Remote / Hybrid',
            employment_type='Internship',
            duration='3_months',
            short_description='Work on cutting-edge generative AI models.',
            description='Detailed responsibilities for AI Engineer.',
            requirements='Python, PyTorch, LLMs',
            is_active=True
        )
        db.session.add(self.job)
        db.session.commit()

        # Create Application (Rahul Kumar)
        self.app_record = JobApplication(
            job_id=self.job.id,
            application_code='AM-APP-001001',
            full_name='Rahul Kumar',
            email='rahul.kumar@example.com',
            phone='+91 9876543210',
            duration='3_months',
            resume_filename='rahul_resume.pdf',
            resume_path='uploads/resumes/rahul_resume.pdf',
            application_fee=399,
            payment_status='paid',
            application_status='submitted',
            application_success_email_status='PENDING'
        )
        db.session.add(self.app_record)
        db.session.commit()

        # Create Employee (AM4827)
        self.employee = Employee(
            employee_id='AM4827',
            application_id=self.app_record.id,
            account_status='active'
        )
        self.employee.set_password('AMTestPass1@')
        db.session.add(self.employee)
        db.session.commit()

    def tearDown(self):
        db.session.remove()
        db.drop_all()
        self.ctx.pop()

    def login_admin(self):
        return self.client.post('/login', data={'email': 'admin@antimatrix.ai', 'password': 'Admin@AntiMatrix2026!'}, follow_redirects=True)

    def test_01_default_template_wording_exact_match(self):
        """Verify the exact wording of the default official templates."""
        # 1. Application Successful Template
        self.assertIn("Application Successfully Submitted — {{Internship Role}} | {{Application ID}}", DEFAULT_APPLICATION_SUCCESSFUL_SUBJECT)
        self.assertIn("Thank you for applying for the **{{Internship Role}} Internship Opportunity at Anti Matrix**.", DEFAULT_APPLICATION_SUCCESSFUL_BODY)
        self.assertIn("### Application Details", DEFAULT_APPLICATION_SUCCESSFUL_BODY)
        self.assertIn("**Application ID:** {{Application ID}}", DEFAULT_APPLICATION_SUCCESSFUL_BODY)
        self.assertIn("**Internship Role:** {{Internship Role}}", DEFAULT_APPLICATION_SUCCESSFUL_BODY)
        self.assertIn("**Application Date:** {{Application Date}}", DEFAULT_APPLICATION_SUCCESSFUL_BODY)
        self.assertIn("Best Regards,\n**Anti Matrix Team**", DEFAULT_APPLICATION_SUCCESSFUL_BODY)

        # 2. Offer Letter / Shortlisted Template
        self.assertIn("Congratulations! You Have Been Shortlisted — {{Internship Role}} | Anti Matrix", DEFAULT_OFFER_LETTER_SUBJECT)
        self.assertIn("**Congratulations! 🎉**", DEFAULT_OFFER_LETTER_BODY)
        self.assertIn("successfully shortlisted for the {{Internship Role}} Internship at Anti Matrix", DEFAULT_OFFER_LETTER_BODY)
        self.assertIn("### Your Internship Details", DEFAULT_OFFER_LETTER_BODY)
        self.assertIn("**Internship Duration:** {{Internship Duration}}", DEFAULT_OFFER_LETTER_BODY)
        self.assertIn("**Start Date:** {{Start Date}}", DEFAULT_OFFER_LETTER_BODY)
        self.assertIn("welcome to the Anti Matrix internship program!** 🚀", DEFAULT_OFFER_LETTER_BODY)

    def test_02_markdown_to_html_email_rendering(self):
        """Verify markdown headings, bold text, emojis, and styling in HTML email."""
        raw_md = "**Congratulations! 🎉**\n\n### Application Details\n\n- Task 1\n- Task 2"
        html_output = markdown_to_html_email(raw_md, title="Test Email")
        
        self.assertIn('<strong style="color: #000000; font-weight: 700;">Congratulations! 🎉</strong>', html_output)
        self.assertIn('<h3 style="font-size: 16px;', html_output)
        self.assertIn('Task 1', html_output)
        self.assertIn('<li', html_output)
        self.assertIn('ANTI<span style="color: #10b981;">-</span>MATRIX', html_output)

    def test_03_application_successful_email_dispatch_and_duplicate_protection(self):
        """Test sending application successful email and verify one-time duplicate protection."""
        self.assertEqual(self.app_record.application_success_email_status, 'PENDING')

        # First send attempt -> should succeed
        success, msg = send_application_successful_email(self.app_record)
        self.assertTrue(success)

        # Check DB state
        updated_app = db.session.get(JobApplication, self.app_record.id)
        self.assertEqual(updated_app.application_success_email_status, 'SENT')
        self.assertIsNotNone(updated_app.application_success_email_sent_at)

        # Check EmailLog
        log = EmailLog.query.filter_by(reference_id=self.app_record.formatted_code).first()
        self.assertIsNotNone(log)
        self.assertEqual(log.recipient_email, 'rahul.kumar@example.com')
        self.assertEqual(log.status, 'SENT')
        self.assertIn('AI Engineer Intern', log.subject)

        # Second send attempt -> MUST be blocked by duplicate protection
        success2, msg2 = send_application_successful_email(self.app_record)
        self.assertFalse(success2)
        self.assertIn('already sent', msg2.lower())

    def test_04_offer_letter_email_dispatch_with_attachment_and_lock(self):
        """Test sending Offer Letter / Shortlisted email with attachment and one-time lock."""
        # First generate Offer Letter DOCX
        emp_doc, path = generate_offer_letter_docx(self.employee)
        self.assertTrue(os.path.exists(path))
        self.assertEqual(emp_doc.email_status, 'not_sent')

        # Send Shortlisted Offer Letter email
        success, msg = send_offer_letter_shortlisted_email(self.employee, start_date='15/09/2026')
        self.assertTrue(success)

        # Check DB state
        updated_doc = db.session.get(EmployeeDocument, emp_doc.id)
        self.assertEqual(updated_doc.email_status, 'sent')
        self.assertEqual(updated_doc.status, 'SENT')
        self.assertIsNotNone(updated_doc.sent_at)

        # Check EmailLog
        log = EmailLog.query.filter_by(reference_id=self.employee.employee_id).first()
        self.assertIsNotNone(log)
        self.assertEqual(log.recipient_email, 'rahul.kumar@example.com')
        self.assertEqual(log.status, 'SENT')
        self.assertTrue(log.has_attachment)
        self.assertIn('Offer_Letter.docx', log.attachment_name)
        self.assertIn('Congratulations! You Have Been Shortlisted', log.subject)

        # Second send attempt -> MUST be blocked
        success2, msg2 = send_offer_letter_shortlisted_email(self.employee)
        self.assertFalse(success2)
        self.assertIn('already sent', msg2.lower())

    def test_05_admin_email_template_preview_endpoint(self):
        """Test preview endpoint with sample preview data."""
        self.login_admin()

        # Preview Application Successful
        res = self.client.get('/admin/templates/email/application_successful/preview')
        self.assertEqual(res.status_code, 200)
        data = res.get_json()
        self.assertIn('Application Successfully Submitted — AI Engineer Intern | AM-APP-1001', data['subject'])
        self.assertIn('Rahul Kumar', data['body_text'])
        self.assertIn('info@antimatrix.co.in', data['body_text'])
        self.assertIn('www.antimatrix.co.in', data['body_text'])

        # Preview Offer Letter / Shortlisted
        res2 = self.client.get('/admin/templates/email/offer_letter/preview')
        self.assertEqual(res2.status_code, 200)
        data2 = res2.get_json()
        self.assertIn('Congratulations! You Have Been Shortlisted — AI Engineer Intern | Anti Matrix', data2['subject'])
        self.assertIn('Rahul Kumar', data2['body_text'])
        self.assertIn('3 Months', data2['body_text'])

    def test_06_admin_send_test_email(self):
        """Test sending a test email without affecting live candidate records."""
        self.login_admin()

        res = self.client.post('/admin/templates/email/offer_letter/test', data={
            'test_recipient': 'test.inbox@example.com'
        }, follow_redirects=True)
        self.assertEqual(res.status_code, 200)

        # Verify test log created
        log = EmailLog.query.filter_by(recipient_email='test.inbox@example.com', template_type='test').first()
        self.assertIsNotNone(log)
        self.assertEqual(log.status, 'SENT')
        self.assertIn('[TEST PREVIEW]', log.subject)

        # Verify live employee record is UNTOUCHED
        emp_doc = self.employee.offer_letter_doc
        self.assertIsNone(emp_doc)

    def test_07_cashfree_return_triggers_application_success_email(self):
        """Verify that successful payment return marks payment paid and status APPLIED."""
        # Create a pending payment
        from models import Payment
        payment = Payment(
            application_id=self.app_record.id,
            cashfree_order_id='order_test_email_trigger_123',
            amount=399.0,
            currency='INR',
            payment_status='pending'
        )
        db.session.add(payment)
        db.session.commit()

        # Simulate Cashfree Return SUCCESS
        res = self.client.get('/payment/cashfree/return?order_id=order_test_email_trigger_123&sim_status=SUCCESS', follow_redirects=True)
        self.assertEqual(res.status_code, 200)

        # Verify application status and email status
        db.session.expire_all()
        updated_app = db.session.get(JobApplication, self.app_record.id)
        self.assertEqual(updated_app.payment_status, 'paid')
        self.assertEqual(updated_app.status, 'APPLIED')
        self.assertEqual(updated_app.application_success_email_status, 'PENDING')

    def test_08_get_request_on_send_application_email_redirects_safely(self):
        """GET request on send-application-email should safely redirect without 500 error."""
        self.login_admin()
        res = self.client.get(f'/admin/applications/{self.app_record.id}/send-application-email', follow_redirects=True)
        self.assertEqual(res.status_code, 200)
        self.assertIn(b'Rahul Kumar', res.data)

    def test_09_nonexistent_application_send_email_graceful(self):
        """Non-existent application ID lookup returns friendly danger message rather than 500."""
        self.login_admin()
        res = self.client.post('/admin/applications/999999/send-application-email', follow_redirects=True)
        self.assertEqual(res.status_code, 200)
        self.assertIn(b'record not found', res.data)

    def test_10_missing_candidate_email_validation(self):
        """Candidate with missing email address yields friendly error message rather than HTTP 500."""
        no_email_app = JobApplication(
            job_id=self.job.id,
            application_code='AM-APP-999999',
            full_name='No Email Candidate',
            email='',
            phone='+91 9999999999',
            resume_filename='test.pdf',
            resume_path='uploads/resumes/test.pdf',
            payment_status='paid',
            status='UNDER_REVIEW',
            application_status='UNDER_REVIEW'
        )
        db.session.add(no_email_app)
        db.session.commit()

        self.login_admin()
        res = self.client.post(f'/admin/applications/{no_email_app.id}/send-application-email', follow_redirects=True)
        self.assertEqual(res.status_code, 200)
        self.assertIn(b'email address is missing', res.data)

        # Status should NOT be marked SENT
        db.session.expire_all()
        updated = db.session.get(JobApplication, no_email_app.id)
        self.assertNotEqual(updated.application_success_email_status, 'SENT')

    def test_11_send_application_email_preserves_application_stage_under_review(self):
        """Sending Application Successful email must NOT prematurely advance status to SHORTLISTED."""
        self.app_record.status = 'UNDER_REVIEW'
        self.app_record.application_status = 'UNDER_REVIEW'
        db.session.commit()

        self.login_admin()
        res = self.client.post(f'/admin/applications/{self.app_record.id}/send-application-email', follow_redirects=True)
        self.assertEqual(res.status_code, 200)

        db.session.expire_all()
        updated_app = db.session.get(JobApplication, self.app_record.id)
        self.assertEqual(updated_app.application_success_email_status, 'SENT')
        # Crucial check: application status remains UNDER_REVIEW!
        self.assertEqual(updated_app.status, 'UNDER_REVIEW')
        self.assertEqual(updated_app.application_status, 'UNDER_REVIEW')

    def test_12_smtp_authentication_failure_gracefully_handled(self):
        """Invalid SMTP credentials should set FAILED status with clean message, never uncaught 500."""
        import os
        from unittest.mock import patch
        import smtplib

        self.app_record.application_success_email_status = 'PENDING'
        db.session.commit()

        self.login_admin()

        with patch('smtplib.SMTP') as mock_smtp:
            instance = mock_smtp.return_value.__enter__.return_value
            instance.login.side_effect = smtplib.SMTPAuthenticationError(535, b'5.7.8 Authentication failed')

            # Set SMTP env vars so send_mime_email attempts SMTP
            with patch.dict(os.environ, {
                'BREVO_API_KEY': '',
                'SMTP_SERVER': 'smtp-relay.brevo.com',
                'SMTP_PORT': '587',
                'SMTP_USER': 'bad_user',
                'SMTP_PASSWORD': 'bad_password',
                'SENDER_EMAIL': 'info@antimatrix.co.in'
            }):
                res = self.client.post(f'/admin/applications/{self.app_record.id}/send-application-email', follow_redirects=True)
                self.assertEqual(res.status_code, 200)
                self.assertIn(b'Email authentication failed', res.data)

        db.session.expire_all()
        updated = db.session.get(JobApplication, self.app_record.id)
        self.assertEqual(updated.application_success_email_status, 'FAILED')

    def test_13_database_records_and_ids_preserved(self):
        """Verify that sending emails does not modify or delete existing job postings or users."""
        job_count_before = JobPosting.query.count()
        user_count_before = User.query.count()
        app_id_before = self.app_record.id
        app_code_before = self.app_record.formatted_code

        self.login_admin()
        self.client.post(f'/admin/applications/{self.app_record.id}/send-application-email', follow_redirects=True)

        self.assertEqual(JobPosting.query.count(), job_count_before)
        self.assertEqual(User.query.count(), user_count_before)
        self.assertEqual(self.app_record.id, app_id_before)
        self.assertEqual(self.app_record.formatted_code, app_code_before)


if __name__ == '__main__':
    unittest.main()
