import os
import io
import unittest
from datetime import datetime, timezone
import docx
from flask import session
from app import create_app
from models import (
    db, User, JobPosting, JobApplication, Employee,
    DocumentTemplate, EmailTemplate, EmployeeDocument, EmailLog
)
from services.offer_letter_service import generate_offer_letter_docx
from services.email_service import (
    render_application_successful_email,
    render_shortlisted_offer_email,
    render_joining_credentials_email,
    send_application_successful_email,
    send_offer_letter_shortlisted_email,
    send_joining_credentials_email
)


class TestAdminCandidateWorkflow(unittest.TestCase):
    """
    Comprehensive Test Suite for the 5-Stage Recruitment and Onboarding Workflow:
    Stage 1: APPLIED
    Stage 2: UNDER REVIEW
    Stage 3: SHORTLISTED (Creates Employee ID, Hashes Temp Password, Clones Offer Letter DOCX)
    Stage 4: OFFER COMPLETED (Dispatches Offer Letter Email + DOCX Attachment via Brevo)
    Stage 5: HIRED (Confirms Joining Date, Dispatches Joining & Employee Credentials Email via Brevo)
    + Student Login via Employee ID (AM...) or Email
    + Template Management for Joining & Employee Credentials
    """

    def setUp(self):
        self.app = create_app('testing')
        self.app_context = self.app.app_context()
        self.app_context.push()
        self.client = self.app.test_client()
        db.create_all()

        # Create Admin Account
        self.admin = User(
            name='Test Administrator',
            email='admin@antimatrix.co.in',
            role='admin',
            is_active=True
        )
        self.admin.set_password('Admin@123456')
        db.session.add(self.admin)

        # Create Candidate Account
        self.candidate = User(
            name='Praveen R',
            email='candidate@example.com',
            role='student',
            is_active=True
        )
        self.candidate.set_password('Student@123456')
        db.session.add(self.candidate)

        # Create Job Posting
        self.job = JobPosting(
            title='AI Engineer Intern',
            department='Artificial Intelligence',
            location='Remote',
            employment_type='Internship',
            duration='3_months',
            short_description='Develop production machine learning models and NLP pipelines.',
            description='Detailed description for AI Engineer Intern.',
            skills='Python, PyTorch, Transformers, FastAPI',
            is_active=True
        )
        db.session.add(self.job)

        # Create Master Offer Letter Template DOCX
        self.templates_dir = os.path.join(self.app.root_path, 'uploads', 'templates')
        os.makedirs(self.templates_dir, exist_ok=True)
        self.master_docx_path = os.path.join(self.templates_dir, 'master_offer_test.docx')

        doc = docx.Document()
        doc.add_heading('ANTI-MATRIX OFFER OF INTERNSHIP', 0)
        doc.add_paragraph('Date: [DD/MM/YYYY]')
        doc.add_paragraph('Dear [Candidate Name],')
        doc.add_paragraph('We are pleased to offer you the position of [Job Title] (Ref: [Reference Number]) in the {{department}} department.')
        doc.add_paragraph('Internship Duration: {{internship_duration}}.')
        doc.add_paragraph('Responsibilities: [brief description of responsibilities]')
        doc.add_paragraph('Joining Date: [Joining Date]. Work Mode: [remote / hybrid / on-site].')
        doc.add_paragraph('Acceptance Deadline: [Acceptance Deadline].')
        doc.save(self.master_docx_path)

        self.doc_template = DocumentTemplate(
            template_type='offer_letter',
            name='Anti-Matrix Master Offer Letter',
            filename='master_offer_test.docx',
            file_path=self.master_docx_path,
            is_active=True
        )
        db.session.add(self.doc_template)

        # Seed Standard Email Templates
        self.app_email_tmpl = EmailTemplate(
            template_type='application_successful',
            name='Application Successful Confirmation',
            subject='Application Successfully Submitted — {{Internship Role}} | {{Application ID}}',
            body="""Dear {{Student Name}},

Thank you for applying for the **{{Internship Role}} Internship Opportunity at Anti Matrix**.
Your Application ID is **{{Application ID}}** submitted on **{{Application Date}}**.
Contact us at {{Company Email}}."""
        )
        db.session.add(self.app_email_tmpl)

        self.offer_email_tmpl = EmailTemplate(
            template_type='offer_letter',
            name='Offer Letter Delivery',
            subject='Congratulations! You Have Been Shortlisted — {{Internship Role}} | Anti Matrix',
            body="""Dear {{Student Name}},

Congratulations! You have been shortlisted for {{Internship Role}}.
Application ID: {{Application ID}}
Duration: {{Internship Duration}}
Start Date: {{Start Date}}."""
        )
        db.session.add(self.offer_email_tmpl)

        self.joining_email_tmpl = EmailTemplate(
            template_type='joining_credentials',
            name='Joining & Employee Credentials',
            subject='Welcome to Anti Matrix — Your Internship Joining Details & Employee Credentials',
            body="""Dear {{Student Name}},

Welcome to the Anti Matrix internship program.
Employee ID: {{Employee ID}}
Temporary Password: {{Employee Password}}
Portal: {{Internship Portal Link}}
Joining Date: {{Joining Date}}."""
        )
        db.session.add(self.joining_email_tmpl)

        db.session.commit()

    def tearDown(self):
        db.session.remove()
        db.drop_all()
        self.app_context.pop()
        if os.path.exists(self.master_docx_path):
            try:
                os.remove(self.master_docx_path)
            except Exception:
                pass

    def login_admin(self):
        return self.client.post('/login', data={
            'email': 'admin@antimatrix.co.in',
            'password': 'Admin@123456'
        }, follow_redirects=True)

    def login_candidate(self):
        return self.client.post('/login', data={
            'email': 'candidate@example.com',
            'password': 'Student@123456'
        }, follow_redirects=True)

    def test_01_candidate_application_submission_defaults_to_applied(self):
        """Test that submitting an application sets status=APPLIED and does NOT auto-send email immediately."""
        self.login_candidate()

        # Submit Application
        resume_bytes = io.BytesIO(b"%PDF-1.4 mock resume binary content")
        aadhaar_bytes = io.BytesIO(b"%PDF-1.4 mock aadhaar binary content")
        data = {
            'first_name': 'Praveen',
            'last_name': 'R',
            'email': 'candidate@example.com',
            'phone': '9876543210',
            'address': '123 Tech Park',
            'state': 'Tamil Nadu',
            'city': 'Chennai',
            'pincode': '600001',
            'education_level': 'Undergraduate',
            'degree': 'B.Tech',
            'major': 'Computer Science',
            'college': 'Anna University',
            'department': 'Computer Science',
            'year_of_study': 'Final Year',
            'graduation_year': '2026',
            'current_cgpa': '9.2',
            'resume': (resume_bytes, 'praveen_resume.pdf'),
            'aadhaar': (aadhaar_bytes, 'praveen_aadhaar.pdf')
        }
        res = self.client.post(f'/careers/apply/{self.job.id}', data=data, content_type='multipart/form-data', follow_redirects=True)
        self.assertEqual(res.status_code, 200)

        app_record = JobApplication.query.filter_by(email='candidate@example.com').first()
        self.assertIsNotNone(app_record)

        # Complete Test Payment
        pay_res = self.client.post(f'/careers/apply/test-payment/{app_record.id}', follow_redirects=True)
        self.assertEqual(pay_res.status_code, 200)

        updated_app = db.session.get(JobApplication, app_record.id)
        self.assertEqual(updated_app.status, 'APPLIED')
        self.assertEqual(updated_app.application_status, 'APPLIED')
        self.assertEqual(updated_app.status_display, 'Applied')

        # Application Successful Email must NOT be sent automatically (status remains PENDING)
        self.assertEqual(updated_app.application_success_email_status, 'PENDING')
        self.assertIsNone(updated_app.application_success_email_sent_at)

        # Candidate My Applications must show 'Applied'
        my_apps_res = self.client.get('/my-applications')
        self.assertIn(b'Applied', my_apps_res.data)
        self.assertIn(updated_app.formatted_code.encode(), my_apps_res.data)

    def test_02_stage_transitions_under_review_and_email(self):
        """Test Admin detail view in APPLIED stage and transition to UNDER_REVIEW."""
        app_record = JobApplication(
            job_id=self.job.id,
            user_id=self.candidate.id,
            full_name='Praveen R',
            email='candidate@example.com',
            phone='9876543210',
            duration='3_months',
            application_fee=399,
            payment_status='paid',
            status='APPLIED',
            application_status='APPLIED',
            resume_filename='resume.pdf',
            resume_path='/fake/resume.pdf'
        )
        db.session.add(app_record)
        db.session.commit()

        self.login_admin()

        # Admin opens application detail page
        detail_res = self.client.get(f'/admin/applications/{app_record.id}')
        self.assertEqual(detail_res.status_code, 200)
        self.assertIn(b'Mark as Under Review', detail_res.data)

        # Admin clicks Mark as Under Review
        post_res = self.client.post(f'/admin/applications/{app_record.id}/mark-under-review', follow_redirects=True)
        self.assertEqual(post_res.status_code, 200)

        updated_app = db.session.get(JobApplication, app_record.id)
        self.assertEqual(updated_app.status, 'UNDER_REVIEW')
        self.assertEqual(updated_app.status_display, 'Under Review')

        # Send Application Successful Email via Brevo
        send_res = self.client.post(f'/admin/applications/{app_record.id}/send-application-email', follow_redirects=True)
        self.assertEqual(send_res.status_code, 200)

        updated_app = db.session.get(JobApplication, app_record.id)
        self.assertEqual(updated_app.application_success_email_status, 'SENT')
        self.assertIsNotNone(updated_app.application_success_email_sent_at)

    def test_03_shortlisting_creates_employee_and_generates_docx(self):
        """Test marking application as SHORTLISTED creates Employee record with random AM ID, temp password, and clones Offer Letter DOCX."""
        app_record = JobApplication(
            job_id=self.job.id,
            user_id=self.candidate.id,
            full_name='Praveen R',
            email='candidate@example.com',
            phone='9876543210',
            duration='3_months',
            application_fee=399,
            payment_status='paid',
            status='UNDER_REVIEW',
            application_status='UNDER_REVIEW',
            resume_filename='resume.pdf',
            resume_path='/fake/resume.pdf'
        )
        db.session.add(app_record)
        db.session.commit()

        self.login_admin()

        # Admin clicks Mark as Shortlisted
        post_res = self.client.post(f'/admin/applications/{app_record.id}/mark-shortlisted', follow_redirects=True)
        self.assertEqual(post_res.status_code, 200)

        updated_app = db.session.get(JobApplication, app_record.id)
        self.assertEqual(updated_app.status, 'SHORTLISTED')
        self.assertEqual(updated_app.status_display, 'Shortlisted')

        # Verify Employee record was created automatically
        self.assertIsNotNone(updated_app.employee)
        emp = updated_app.employee
        self.assertTrue(emp.employee_id.startswith('AM'))
        self.assertEqual(len(emp.employee_id), 6)  # AM + 4 digits
        self.assertTrue(emp.password_hash.startswith(('scrypt:', 'pbkdf2:')))

        # Verify Offer Letter DOCX was generated automatically
        offer_doc = updated_app.offer_letter_doc
        self.assertIsNotNone(offer_doc)
        self.assertTrue(os.path.exists(offer_doc.file_path))

        # Check DOCX contents
        gen_doc = docx.Document(offer_doc.file_path)
        full_text = "\n".join([p.text for p in gen_doc.paragraphs])
        self.assertIn('Praveen R', full_text)
        self.assertIn(updated_app.formatted_code, full_text)
        self.assertIn('AI Engineer Intern', full_text)

    def test_04_mark_complete_dispatches_offer_letter(self):
        """Test Stage 4: Mark as Complete sets status to OFFER_COMPLETED and dispatches Offer Letter email."""
        app_record = JobApplication(
            job_id=self.job.id,
            user_id=self.candidate.id,
            full_name='Praveen R',
            email='candidate@example.com',
            phone='9876543210',
            duration='3_months',
            application_fee=399,
            payment_status='paid',
            status='SHORTLISTED',
            application_status='SHORTLISTED',
            resume_filename='resume.pdf',
            resume_path='/fake/resume.pdf'
        )
        db.session.add(app_record)
        db.session.commit()

        # Shortlist candidate (creates employee & offer doc)
        self.login_admin()
        self.client.post(f'/admin/applications/{app_record.id}/mark-shortlisted', follow_redirects=True)

        # Admin clicks Mark as Complete
        complete_res = self.client.post(f'/admin/applications/{app_record.id}/mark-complete', follow_redirects=True)
        self.assertEqual(complete_res.status_code, 200)

        updated_app = db.session.get(JobApplication, app_record.id)
        self.assertEqual(updated_app.status, 'OFFER_COMPLETED')
        self.assertEqual(updated_app.status_display, 'Offer Completed')
        self.assertIsNotNone(updated_app.offer_completed_at)

        offer_doc = updated_app.offer_letter_doc
        self.assertEqual(offer_doc.email_status, 'sent')
        self.assertIsNotNone(offer_doc.sent_at)

    def test_05_mark_hired_sets_joining_date_and_dispatches_credentials_email(self):
        """Test Stage 5: Mark as Hired sets status to HIRED, updates joining_date, and dispatches Joining Credentials email."""
        app_record = JobApplication(
            job_id=self.job.id,
            user_id=self.candidate.id,
            full_name='Praveen R',
            email='candidate@example.com',
            phone='9876543210',
            duration='3_months',
            application_fee=399,
            payment_status='paid',
            status='OFFER_COMPLETED',
            application_status='OFFER_COMPLETED',
            resume_filename='resume.pdf',
            resume_path='/fake/resume.pdf'
        )
        db.session.add(app_record)
        db.session.commit()

        self.login_admin()
        # Ensure employee and offer doc exist
        self.client.post(f'/admin/applications/{app_record.id}/mark-shortlisted', follow_redirects=True)
        self.client.post(f'/admin/applications/{app_record.id}/mark-complete', follow_redirects=True)

        # Admin submits Mark as Hired with joining_date
        hired_res = self.client.post(f'/admin/applications/{app_record.id}/mark-hired', data={
            'joining_date': '15/09/2026'
        }, follow_redirects=True)
        self.assertEqual(hired_res.status_code, 200)

        updated_app = db.session.get(JobApplication, app_record.id)
        self.assertEqual(updated_app.status, 'HIRED')
        self.assertEqual(updated_app.status_display, 'Hired')
        self.assertEqual(updated_app.joining_date, '15/09/2026')
        self.assertIsNotNone(updated_app.hired_at)
        self.assertEqual(updated_app.joining_email_status, 'SENT')
        self.assertIsNotNone(updated_app.joining_email_sent_at)

        # Encrypted temp password must be purged after sending joining email
        self.assertIsNone(updated_app.employee.temp_password_encrypted)

    def test_06_student_login_with_employee_id_and_email(self):
        """Test candidate login using both Employee ID (AM...) and registered Email."""
        emp_id = Employee.generate_unique_employee_id()
        plaintext_pwd = 'AMx7P9k2Q!#1'

        app_record = JobApplication(
            job_id=self.job.id,
            user_id=self.candidate.id,
            full_name='Praveen R',
            email='candidate@example.com',
            phone='9876543210',
            duration='3_months',
            application_fee=399,
            payment_status='paid',
            status='HIRED',
            application_status='HIRED',
            resume_filename='resume.pdf',
            resume_path='/fake/resume.pdf'
        )
        db.session.add(app_record)
        db.session.flush()

        employee = Employee(
            employee_id=emp_id,
            application_id=app_record.id,
            account_status='active'
        )
        employee.set_password(plaintext_pwd)
        db.session.add(employee)
        db.session.commit()

        # 1. Login with Employee ID (AM...) + Temporary Password
        emp_login_res = self.client.post('/login', data={
            'email': emp_id,
            'password': plaintext_pwd
        }, follow_redirects=True)
        self.assertEqual(emp_login_res.status_code, 200)

        # Verify authenticated session
        dashboard_res = self.client.get('/my-applications')
        self.assertEqual(dashboard_res.status_code, 200)
        self.assertIn(b'Praveen', dashboard_res.data)

        self.client.get('/logout')

        # 2. Login with registered Email + Temporary Password
        email_login_res = self.client.post('/login', data={
            'email': 'candidate@example.com',
            'password': plaintext_pwd
        }, follow_redirects=True)
        self.assertEqual(email_login_res.status_code, 200)

        # 3. Invalid credentials test
        self.client.get('/logout')
        bad_login_res = self.client.post('/login', data={
            'email': emp_id,
            'password': 'WrongPassword123'
        }, follow_redirects=True)
        self.assertIn(b'Invalid Employee ID or password', bad_login_res.data)

    def test_07_template_management_joining_credentials(self):
        """Test template management for joining_credentials template type."""
        self.login_admin()

        # Update Joining Credentials Template
        update_res = self.client.post('/admin/templates/email/joining_credentials', data={
            'subject': 'Custom Joining Subject: {{Student Name}} | {{Employee ID}}',
            'body': 'Welcome {{Student Name}}! Your ID is {{Employee ID}} and joining date is {{Joining Date}}.'
        }, follow_redirects=True)
        self.assertEqual(update_res.status_code, 200)

        tmpl = EmailTemplate.query.filter_by(template_type='joining_credentials').first()
        self.assertIsNotNone(tmpl)
        self.assertEqual(tmpl.subject, 'Custom Joining Subject: {{Student Name}} | {{Employee ID}}')

        # Preview Joining Credentials Template API
        prev_res = self.client.get('/admin/templates/email/joining_credentials/preview')
        self.assertEqual(prev_res.status_code, 200)
        prev_data = prev_res.get_json()
        self.assertIn('Rahul Kumar', prev_data['subject'])
        self.assertIn('AM4827', prev_data['subject'])


if __name__ == '__main__':
    unittest.main()


if __name__ == '__main__':
    unittest.main()
