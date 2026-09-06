import unittest
import os
import io
import shutil
import tempfile
import docx
from datetime import datetime, timezone
from app import create_app
from models import db, User, JobPosting, JobApplication, Employee, EmployeeDocument, DocumentTemplate, EmailTemplate, Payment, EmailLog
from services.offer_letter_service import (
    generate_offer_letter_docx, send_offer_letter_email,
    get_active_offer_letter_template, determine_job_category,
    ensure_default_templates_initialized, OFFER_LETTER_CATEGORIES,
    OfferLetterTemplateNotFoundError, OfferLetterTemplateFileMissingError
)


class TestOfferLetterSystem(unittest.TestCase):
    def setUp(self):
        """Set up test Flask app with in-memory or test SQLite database."""
        self.app = create_app('testing')
        self.app.config['WTF_CSRF_ENABLED'] = False
        
        self.client = self.app.test_client()
        self.ctx = self.app.app_context()
        self.ctx.push()
        
        db.create_all()

        # Ensure templates directory exists
        self.templates_dir = os.path.join(self.app.root_path, 'uploads', 'templates')
        os.makedirs(self.templates_dir, exist_ok=True)
        
        # Initialize default master templates for all 4 categories
        ensure_default_templates_initialized()

        # Create Job Postings for each of the 4 domains
        self.job_ai = JobPosting(
            title='AI Research & ML Intern',
            department='AI & Data Science',
            location='Remote',
            employment_type='Internship',
            duration='1_month',
            short_description='Develop machine learning pipelines and deep learning algorithms.',
            description='Detailed responsibilities and scope for AI Research & ML Intern.',
            requirements='Python, PyTorch, Scikit-Learn',
            is_active=True
        )

        self.job_web = JobPosting(
            title='Frontend & Backend Web Development Intern',
            department='Engineering',
            location='Remote / Hybrid',
            employment_type='Internship',
            duration='1_month',
            short_description='Build scalable web applications and REST APIs.',
            description='Detailed responsibilities and scope for Web Development Intern.',
            requirements='HTML, CSS, JavaScript, React, Flask',
            is_active=True
        )

        self.job_app = JobPosting(
            title='Mobile App Development Intern (Flutter / Android)',
            department='Mobile Engineering',
            location='Remote',
            employment_type='Internship',
            duration='3_months',
            short_description='Design and build responsive mobile applications.',
            description='Detailed responsibilities and scope for Mobile App Development Intern.',
            requirements='Flutter, Dart, Android SDK',
            is_active=True
        )

        self.job_data = JobPosting(
            title='Data Analytics & Business Intelligence Intern',
            department='Analytics',
            location='Remote',
            employment_type='Internship',
            duration='1_month',
            short_description='Analyze datasets and create business dashboards.',
            description='Detailed responsibilities and scope for Data Analytics Intern.',
            requirements='SQL, Python, Power BI, Excel',
            is_active=True
        )

        self.job_unsupported = JobPosting(
            title='Executive Chef Consultant',
            department='Culinary',
            location='On-site',
            employment_type='Contract',
            duration='6_months',
            short_description='Culinary planning.',
            description='Detailed responsibilities for Executive Chef Consultant.',
            requirements='Culinary Arts',
            is_active=True
        )

        db.session.add_all([self.job_ai, self.job_web, self.job_app, self.job_data, self.job_unsupported])
        db.session.commit()

        # Create Candidate Applications
        self.app_ai = JobApplication(
            job_id=self.job_ai.id,
            application_code='AM-APP-AI001',
            full_name='Aarav Sharma',
            email='aarav.sharma@example.com',
            phone='+91 9876543211',
            duration='1_month',
            resume_filename='aarav_resume.pdf',
            resume_path='uploads/resumes/aarav_resume.pdf',
            application_fee=399,
            payment_status='paid',
            application_status='shortlisted'
        )

        self.app_web = JobApplication(
            job_id=self.job_web.id,
            application_code='AM-APP-WEB002',
            full_name='Bhavna Patel',
            email='bhavna.patel@example.com',
            phone='+91 9876543212',
            duration='1_month',
            resume_filename='bhavna_resume.pdf',
            resume_path='uploads/resumes/bhavna_resume.pdf',
            application_fee=399,
            payment_status='paid',
            application_status='shortlisted'
        )

        self.app_app = JobApplication(
            job_id=self.job_app.id,
            application_code='AM-APP-APP003',
            full_name='Chetan Verma',
            email='chetan.verma@example.com',
            phone='+91 9876543213',
            duration='3_months',
            resume_filename='chetan_resume.pdf',
            resume_path='uploads/resumes/chetan_resume.pdf',
            application_fee=399,
            payment_status='paid',
            application_status='shortlisted'
        )

        self.app_data = JobApplication(
            job_id=self.job_data.id,
            application_code='AM-APP-DATA004',
            full_name='Divya Nair',
            email='divya.nair@example.com',
            phone='+91 9876543214',
            duration='1_month',
            resume_filename='divya_resume.pdf',
            resume_path='uploads/resumes/divya_resume.pdf',
            application_fee=399,
            payment_status='paid',
            application_status='shortlisted'
        )

        self.app_unsupported = JobApplication(
            job_id=self.job_unsupported.id,
            application_code='AM-APP-CHEF005',
            full_name='Eshan Rao',
            email='eshan.rao@example.com',
            phone='+91 9876543215',
            duration='6_months',
            resume_filename='eshan_resume.pdf',
            resume_path='uploads/resumes/eshan_resume.pdf',
            application_fee=399,
            payment_status='paid',
            application_status='shortlisted'
        )

        db.session.add_all([self.app_ai, self.app_web, self.app_app, self.app_data, self.app_unsupported])
        db.session.commit()

        # Create Employees
        self.emp_ai = Employee(employee_id='AMAI101', application_id=self.app_ai.id, account_status='active')
        self.emp_ai.set_password('EmpPass123@')

        self.emp_web = Employee(employee_id='AMWEB102', application_id=self.app_web.id, account_status='active')
        self.emp_web.set_password('EmpPass123@')

        self.emp_app = Employee(employee_id='AMAPP103', application_id=self.app_app.id, account_status='active')
        self.emp_app.set_password('EmpPass123@')

        self.emp_data = Employee(employee_id='AMDATA104', application_id=self.app_data.id, account_status='active')
        self.emp_data.set_password('EmpPass123@')

        self.emp_unsupported = Employee(employee_id='AMCHEF105', application_id=self.app_unsupported.id, account_status='active')
        self.emp_unsupported.set_password('EmpPass123@')

        db.session.add_all([self.emp_ai, self.emp_web, self.emp_app, self.emp_data, self.emp_unsupported])

        # Create Admin & Member users if not exist
        self.admin = User.query.filter_by(email='admin@antimatrix.ai').first()
        if not self.admin:
            self.admin = User(name='Admin User', email='admin@antimatrix.ai', role='admin', is_active=True)
            self.admin.set_password('Admin@AntiMatrix2026!')
            db.session.add(self.admin)
        else:
            self.admin.role = 'admin'
            self.admin.set_password('Admin@AntiMatrix2026!')

        self.member = User.query.filter_by(email='member@example.com').first()
        if not self.member:
            self.member = User(name='Member User', email='member@example.com', role='member', is_active=True)
            self.member.set_password('Member@2026!')
            db.session.add(self.member)

        db.session.commit()

    def login_admin(self):
        return self.client.post('/login', data={'email': 'admin@antimatrix.ai', 'password': 'Admin@AntiMatrix2026!'}, follow_redirects=True)

    def login_member(self):
        return self.client.post('/login', data={'email': 'member@example.com', 'password': 'Member@2026!'}, follow_redirects=True)

    def logout(self):
        return self.client.get('/logout', follow_redirects=True)

    def tearDown(self):
        db.session.remove()
        db.drop_all()
        self.ctx.pop()

    def test_01_job_category_resolution(self):
        """Test intelligent determination of internship category from job titles and descriptions."""
        self.assertEqual(determine_job_category(self.job_ai), 'AI_ML')
        self.assertEqual(determine_job_category(self.job_web), 'WEB_DEVELOPMENT')
        self.assertEqual(determine_job_category(self.job_app), 'APP_DEVELOPMENT')
        self.assertEqual(determine_job_category(self.job_data), 'DATA_ANALYTICS')
        self.assertIsNone(determine_job_category(self.job_unsupported))

        # Test string inputs
        self.assertEqual(determine_job_category('AI Research Intern'), 'AI_ML')
        self.assertEqual(determine_job_category('Full Stack Web Developer'), 'WEB_DEVELOPMENT')
        self.assertEqual(determine_job_category('Flutter Android App Developer'), 'APP_DEVELOPMENT')
        self.assertEqual(determine_job_category('Power BI Data Analyst'), 'DATA_ANALYTICS')
        self.assertIsNone(determine_job_category('Marketing Sales Executive'))

    def test_02_generate_ai_ml_offer_letter(self):
        """Test AI & ML offer letter generation: uses AI & ML template, role wording, and candidate details."""
        emp_doc, output_path = generate_offer_letter_docx(self.emp_ai)
        self.assertTrue(os.path.exists(output_path))
        self.assertEqual(emp_doc.employee_id, self.emp_ai.id)

        # Inspect generated DOCX content
        doc = docx.Document(output_path)
        doc_text = "\n".join([p.text for p in doc.paragraphs])

        self.assertIn('Aarav Sharma', doc_text)
        self.assertIn(self.app_ai.formatted_code, doc_text)
        self.assertIn('AI & ML Intern', doc_text)
        self.assertIn('Artificial Intelligence and Machine Learning', doc_text)
        self.assertIn('data processing, model development', doc_text)
        self.assertIn('Anti-Matrix', doc_text)
        self.assertNotIn('[Candidate Name]', doc_text)
        self.assertNotIn('[Reference Number]', doc_text)

    def test_03_generate_web_dev_offer_letter(self):
        """Test Web Development offer letter generation: uses Web Dev template, role wording, and candidate details."""
        emp_doc, output_path = generate_offer_letter_docx(self.emp_web)
        self.assertTrue(os.path.exists(output_path))
        self.assertEqual(emp_doc.employee_id, self.emp_web.id)

        doc = docx.Document(output_path)
        doc_text = "\n".join([p.text for p in doc.paragraphs])

        self.assertIn('Bhavna Patel', doc_text)
        self.assertIn(self.app_web.formatted_code, doc_text)
        self.assertIn('Web Development Intern', doc_text)
        self.assertIn('responsive web applications', doc_text)
        self.assertIn('frontend and backend technologies', doc_text)
        self.assertNotIn('[Candidate Name]', doc_text)

    def test_04_generate_app_dev_offer_letter(self):
        """Test App Development offer letter generation: uses App Dev template, role wording, and candidate details."""
        emp_doc, output_path = generate_offer_letter_docx(self.emp_app)
        self.assertTrue(os.path.exists(output_path))
        self.assertEqual(emp_doc.employee_id, self.emp_app.id)

        doc = docx.Document(output_path)
        doc_text = "\n".join([p.text for p in doc.paragraphs])

        self.assertIn('Chetan Verma', doc_text)
        self.assertIn(self.app_app.formatted_code, doc_text)
        self.assertIn('App Development Intern', doc_text)
        self.assertIn('designing and developing mobile applications', doc_text)
        self.assertIn('integrating APIs and databases', doc_text)
        self.assertNotIn('[Candidate Name]', doc_text)

    def test_05_generate_data_analytics_offer_letter(self):
        """Test Data Analytics offer letter generation: uses Data Analytics template, role wording, and candidate details."""
        emp_doc, output_path = generate_offer_letter_docx(self.emp_data)
        self.assertTrue(os.path.exists(output_path))
        self.assertEqual(emp_doc.employee_id, self.emp_data.id)

        doc = docx.Document(output_path)
        doc_text = "\n".join([p.text for p in doc.paragraphs])

        self.assertIn('Divya Nair', doc_text)
        self.assertIn(self.app_data.formatted_code, doc_text)
        self.assertIn('Data Analytics Intern', doc_text)
        self.assertIn('data collection, cleaning, preprocessing', doc_text)
        self.assertIn('visualization, reporting', doc_text)
        self.assertNotIn('[Candidate Name]', doc_text)

    def test_06_unsupported_category_fallback(self):
        """Verify that an unsupported job category raises a clear error and does not generate an incorrect letter."""
        with self.assertRaises(OfferLetterTemplateNotFoundError) as ctx:
            generate_offer_letter_docx(self.emp_unsupported)
        self.assertIn('No job-specific offer letter template is available', str(ctx.exception))

    def test_07_master_templates_never_modified(self):
        """Verify that master templates on disk are never altered during generation."""
        # Generate all 4
        generate_offer_letter_docx(self.emp_ai)
        generate_offer_letter_docx(self.emp_web)
        generate_offer_letter_docx(self.emp_app)
        generate_offer_letter_docx(self.emp_data)

        for cat_key, cat_data in OFFER_LETTER_CATEGORIES.items():
            tmpl = get_active_offer_letter_template(cat_key)
            doc = docx.Document(tmpl.file_path)
            master_text = "\n".join([p.text for p in doc.paragraphs])

            self.assertIn('[DD/MM/YYYY]', master_text)
            self.assertIn('[Candidate Name]', master_text)
            self.assertIn('[Reference Number]', master_text)
            self.assertNotIn('Aarav Sharma', master_text)
            self.assertNotIn('Bhavna Patel', master_text)
            self.assertNotIn('Chetan Verma', master_text)
            self.assertNotIn('Divya Nair', master_text)

    def test_08_replacing_one_template_does_not_affect_others(self):
        """Test that replacing the Web Dev template leaves AI & ML, App Dev, and Data Analytics active and intact."""
        self.login_admin()

        # Capture active IDs before upload
        ai_tmpl_before = get_active_offer_letter_template('AI_ML').id
        web_tmpl_before = get_active_offer_letter_template('WEB_DEVELOPMENT').id
        app_tmpl_before = get_active_offer_letter_template('APP_DEVELOPMENT').id
        data_tmpl_before = get_active_offer_letter_template('DATA_ANALYTICS').id

        # Upload replacement Web Dev template
        new_doc_bytes = io.BytesIO()
        doc = docx.Document()
        doc.add_paragraph("New Custom Web Dev Template V2")
        doc.add_paragraph("Dear [Candidate Name], Ref: [Reference Number], Role: [Job Title]")
        doc.save(new_doc_bytes)
        new_doc_bytes.seek(0)

        res = self.client.post('/admin/templates/document/offer_letter_web_development/upload', data={
            'template_file': (new_doc_bytes, 'new_web_dev_template.docx')
        }, content_type='multipart/form-data', follow_redirects=True)
        self.assertEqual(res.status_code, 200)

        # Verify only Web Dev template ID changed
        ai_tmpl_after = get_active_offer_letter_template('AI_ML').id
        web_tmpl_after = get_active_offer_letter_template('WEB_DEVELOPMENT').id
        app_tmpl_after = get_active_offer_letter_template('APP_DEVELOPMENT').id
        data_tmpl_after = get_active_offer_letter_template('DATA_ANALYTICS').id

        self.assertEqual(ai_tmpl_before, ai_tmpl_after)
        self.assertNotEqual(web_tmpl_before, web_tmpl_after)
        self.assertEqual(app_tmpl_before, app_tmpl_after)
        self.assertEqual(data_tmpl_before, data_tmpl_after)

    def test_09_admin_templates_ui_renders_four_cards(self):
        """Verify the admin templates page renders all 4 separate job-specific offer letter cards."""
        self.login_admin()
        res = self.client.get('/admin/templates')
        self.assertEqual(res.status_code, 200)
        html = res.get_data(as_text=True)

        self.assertIn('AI &amp; ML Internship', html)
        self.assertIn('Web Development', html)
        self.assertIn('App Development', html)
        self.assertIn('Data Analytics', html)
        self.assertIn('offer_letter_ai_ml', html)
        self.assertIn('offer_letter_web_development', html)
        self.assertIn('offer_letter_app_development', html)
        self.assertIn('offer_letter_data_analytics', html)

    def test_10_admin_generate_page_shows_matched_category(self):
        """Verify generate offer letter page displays detected category badge and matching template."""
        self.login_admin()
        res = self.client.get(f'/admin/employees/{self.emp_ai.employee_id}/offer-letter/generate')
        self.assertEqual(res.status_code, 200)
        html = res.get_data(as_text=True)

        self.assertIn('AI &amp; ML Domain', html)
        self.assertIn('AI &amp; ML Internship', html)
        self.assertIn('Generate Offer Letter DOCX', html)


if __name__ == '__main__':
    unittest.main()
