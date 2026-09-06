import os
import shutil
import smtplib
import uuid
import time
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import datetime, timezone, timedelta
import docx
from flask import current_app
from models import db, Employee, JobApplication, JobPosting, EmployeeDocument, DocumentTemplate, EmailTemplate


OFFER_LETTER_CATEGORIES = {
    'AI_ML': {
        'type': 'offer_letter_ai_ml',
        'name': 'AI & ML Internship Offer Letter',
        'card_title': 'AI & ML Internship',
        'default_title': 'AI & ML Intern',
        'default_filename': 'offer_letter_ai_ml_master.docx',
        'badge': 'AI & ML Domain',
        'run0': 'During your internship, you will work on a ',
        'run1': 'real-time project',
        'run2': ' and gain practical experience in Artificial Intelligence and Machine Learning, including data processing, model development, training, testing, evaluation, and implementation.',
        'keywords': [
            'ai & ml', 'ai/ml', 'artificial intelligence', 'machine learning',
            'deep learning', 'data science', 'ai research', 'computer vision',
            'nlp', 'natural language processing', 'ai intern', 'ml intern', 'ai', 'ml'
        ]
    },
    'WEB_DEVELOPMENT': {
        'type': 'offer_letter_web_development',
        'name': 'Web Development Internship Offer Letter',
        'card_title': 'Web Development Internship',
        'default_title': 'Web Development Intern',
        'default_filename': 'offer_letter_web_development_master.docx',
        'badge': 'Web Dev Domain',
        'run0': 'During your internship, you will work on a ',
        'run1': 'real-time web development project',
        'run2': ' and gain practical experience in building responsive web applications, working with frontend and backend technologies, database integration, APIs, testing, and deployment.',
        'keywords': [
            'web development', 'web developer', 'frontend', 'front-end', 'front end',
            'backend', 'back-end', 'back end', 'full stack', 'fullstack', 'full-stack',
            'html', 'css', 'javascript', 'react', 'node', 'vue', 'angular',
            'flask', 'django', 'php', 'web intern', 'web'
        ]
    },
    'APP_DEVELOPMENT': {
        'type': 'offer_letter_app_development',
        'name': 'App Development Internship Offer Letter',
        'card_title': 'App Development Internship',
        'default_title': 'App Development Intern',
        'default_filename': 'offer_letter_app_development_master.docx',
        'badge': 'App Dev Domain',
        'run0': 'During your internship, you will work on a ',
        'run1': 'real-time application development project',
        'run2': ' and gain practical experience in designing and developing mobile applications, implementing user interfaces, integrating APIs and databases, testing, debugging, and preparing applications for deployment.',
        'keywords': [
            'app development', 'app developer', 'mobile application', 'mobile app',
            'mobile developer', 'android', 'ios', 'flutter', 'react native',
            'swift', 'kotlin', 'app intern', 'mobile intern', 'app'
        ]
    },
    'DATA_ANALYTICS': {
        'type': 'offer_letter_data_analytics',
        'name': 'Data Analytics Internship Offer Letter',
        'card_title': 'Data Analytics Internship',
        'default_title': 'Data Analytics Intern',
        'default_filename': 'offer_letter_data_analytics_master.docx',
        'badge': 'Data Analytics Domain',
        'run0': 'During your internship, you will work on a ',
        'run1': 'real-time data analytics project',
        'run2': ' and gain practical experience in data collection, cleaning, preprocessing, analysis, visualization, reporting, and deriving meaningful insights from data.',
        'keywords': [
            'data analytics', 'data analyst', 'business intelligence', 'bi analyst',
            'power bi', 'powerbi', 'tableau', 'data visualization', 'data analysis',
            'sql analyst', 'analytics intern', 'analytics'
        ]
    }
}


def determine_job_category(job_or_title_or_dept):
    """
    Determines the matching internship category key ('AI_ML', 'WEB_DEVELOPMENT', 'APP_DEVELOPMENT', 'DATA_ANALYTICS')
    from a JobPosting object, title string, or department string.
    Returns None if no supported category matches.
    """
    if not job_or_title_or_dept:
        return None

    if isinstance(job_or_title_or_dept, str):
        text = job_or_title_or_dept.strip().lower()
    elif isinstance(job_or_title_or_dept, JobPosting):
        job = job_or_title_or_dept
        # Check explicit category attribute if present
        if hasattr(job, 'category') and job.category:
            cat_upper = str(job.category).upper()
            if cat_upper in OFFER_LETTER_CATEGORIES:
                return cat_upper
        # Combine title, department, skills, short_description
        parts = [
            job.title or '',
            job.department or '',
            getattr(job, 'skills', '') or '',
            getattr(job, 'short_description', '') or ''
        ]
        text = " ".join(parts).lower()
    else:
        text = str(job_or_title_or_dept).lower()

    # Step 1: Check multi-word / specific keyword matches in priority order
    for cat_key in ['DATA_ANALYTICS', 'AI_ML', 'APP_DEVELOPMENT', 'WEB_DEVELOPMENT']:
        cat_info = OFFER_LETTER_CATEGORIES[cat_key]
        for kw in cat_info['keywords']:
            # For multi-word keywords (e.g. "data analytics", "web development", "mobile app")
            if ' ' in kw or '/' in kw or '-' in kw or '&' in kw:
                if kw in text:
                    return cat_key

    # Step 2: Check word tokens / single-word keywords
    words = set(text.replace('/', ' ').replace('-', ' ').replace('&', ' ').replace(',', ' ').split())
    for cat_key in ['DATA_ANALYTICS', 'AI_ML', 'APP_DEVELOPMENT', 'WEB_DEVELOPMENT']:
        cat_info = OFFER_LETTER_CATEGORIES[cat_key]
        for kw in cat_info['keywords']:
            if kw in words or kw in text:
                return cat_key

    return None


def replace_placeholders_in_paragraph(paragraph, mapping):
    """
    Safely replaces placeholder keys with replacement values in a docx Paragraph.
    Preserves exact font, size, bold, italic, color, and run-level styles.
    Handles placeholders both inside individual runs and spanning multiple runs.
    """
    full_text = paragraph.text
    if not full_text:
        return

    # Check if any placeholder exists in the paragraph text
    needs_replacement = False
    for key in mapping:
        if key in full_text:
            needs_replacement = True
            break

    if not needs_replacement:
        return

    for key, val in mapping.items():
        if key not in paragraph.text:
            continue

        str_val = str(val) if val is not None else ''

        # 1. First attempt: check if placeholder is contained entirely in a single run
        replaced_in_single_run = False
        for run in paragraph.runs:
            if key in run.text:
                run.text = run.text.replace(key, str_val)
                replaced_in_single_run = True

        # 2. Fallback: if placeholder spans multiple adjacent runs
        if not replaced_in_single_run and key in paragraph.text:
            combined = "".join([r.text for r in paragraph.runs])
            new_text = combined.replace(key, str_val)
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for r in paragraph.runs[1:]:
                    r.text = ""


class OfferLetterTemplateNotFoundError(Exception):
    """Raised when no active Offer Letter template record exists in the database for the required category."""
    pass


class OfferLetterTemplateFileMissingError(Exception):
    """Raised when the active Offer Letter template record exists in DB, but the physical file is missing from disk/storage."""
    pass


def get_active_offer_letter_template(category_or_job=None):
    """
    Retrieves the active Offer Letter DocumentTemplate record from the database for the specified category or job.
    If category_or_job is None, defaults to 'AI_ML' (with backward compatibility for 'offer_letter').
    Strictly queries the database for an active template and validates that its stored file exists.
    """
    if category_or_job is None:
        cat_key = 'AI_ML'
    elif isinstance(category_or_job, str):
        if category_or_job.upper() in OFFER_LETTER_CATEGORIES:
            cat_key = category_or_job.upper()
        else:
            # Check if it is a template_type string directly (e.g. 'offer_letter_web_development')
            matched_key = None
            for k, info in OFFER_LETTER_CATEGORIES.items():
                if info['type'] == category_or_job:
                    matched_key = k
                    break
            if matched_key:
                cat_key = matched_key
            else:
                cat_key = determine_job_category(category_or_job)
    elif isinstance(category_or_job, JobPosting):
        cat_key = determine_job_category(category_or_job)
    elif isinstance(category_or_job, (JobApplication, Employee)):
        job = category_or_job.job if hasattr(category_or_job, 'job') else None
        cat_key = determine_job_category(job) if job else None
    else:
        cat_key = None

    if not cat_key or cat_key not in OFFER_LETTER_CATEGORIES:
        raise OfferLetterTemplateNotFoundError(
            "No job-specific offer letter template is available for this internship. Please upload the appropriate template before generating the offer letter."
        )

    cat_info = OFFER_LETTER_CATEGORIES[cat_key]
    template_type = cat_info['type']

    # Query active template for this specific category
    active_template = DocumentTemplate.query.filter_by(
        template_type=template_type,
        is_active=True
    ).order_by(DocumentTemplate.id.desc()).first()

    # Backward compatibility for AI_ML category if offer_letter_ai_ml is not yet uploaded
    if not active_template and cat_key == 'AI_ML':
        active_template = DocumentTemplate.query.filter_by(
            template_type='offer_letter',
            is_active=True
        ).order_by(DocumentTemplate.id.desc()).first()

    if not active_template:
        raise OfferLetterTemplateNotFoundError(
            f"No active template found for {cat_info['card_title']}. Please upload the template from Admin Dashboard → Templates."
        )

    if not active_template.file_path or not os.path.exists(active_template.file_path):
        raise OfferLetterTemplateFileMissingError(
            f"The active template file for {cat_info['card_title']} could not be found at '{active_template.file_path}'. Please upload the template again."
        )

    return active_template


def ensure_default_templates_initialized():
    """
    Ensures that default master DOCX files for all 4 internship categories exist in uploads/templates/
    and have corresponding active DocumentTemplate records in the database without altering existing data.
    """
    templates_dir = os.path.join(current_app.root_path, 'uploads', 'templates')
    os.makedirs(templates_dir, exist_ok=True)

    static_defaults_dir = os.path.join(current_app.root_path, 'static', 'default_templates')
    master_ref = os.path.join(templates_dir, 'offer letter (Anti-matrix).docx')
    if not os.path.exists(master_ref):
        alt_ref = os.path.join(templates_dir, 'offer_letter_master.docx')
        if os.path.exists(alt_ref):
            master_ref = alt_ref

    for cat_key, cat_data in OFFER_LETTER_CATEGORIES.items():
        target_path = os.path.join(templates_dir, cat_data['default_filename'])
        static_src = os.path.join(static_defaults_dir, cat_data['default_filename'])
        
        # 1. Create file if missing
        if not os.path.exists(target_path):
            if os.path.exists(static_src):
                shutil.copy2(static_src, target_path)
            elif os.path.exists(master_ref):
                doc = docx.Document(master_ref)
                for p in doc.paragraphs:
                    if 'Congratulations! We are pleased to inform you' in p.text:
                        if len(p.runs) >= 2:
                            p.runs[1].text = cat_data['default_title']
                    elif 'During your internship' in p.text:
                        if len(p.runs) >= 3:
                            p.runs[0].text = cat_data['run0']
                            p.runs[1].text = cat_data['run1']
                            p.runs[2].text = cat_data['run2']
                            for r in p.runs[3:]:
                                r.text = ''
                doc.save(target_path)

        # 2. Check / insert active DocumentTemplate record in DB
        tmpl_record = DocumentTemplate.query.filter_by(
            template_type=cat_data['type'],
            is_active=True
        ).first()

        if not tmpl_record and os.path.exists(target_path):
            new_tmpl = DocumentTemplate(
                template_type=cat_data['type'],
                name=cat_data['name'],
                filename=cat_data['default_filename'],
                file_path=target_path,
                is_active=True
            )
            db.session.add(new_tmpl)

    db.session.commit()


def generate_offer_letter_docx(application_or_employee, custom_params=None, force_regenerate=False):
    """
    Generates a personalized Offer Letter DOCX for the given JobApplication or Employee by cloning
    the category-specific master template.
    Replaces all placeholders while strictly preserving typography, borders, logos, and layout.
    Reuses existing generated file if already generated (idempotent) unless force_regenerate=True.
    """
    if isinstance(application_or_employee, Employee):
        employee = application_or_employee
        app = employee.application
    elif isinstance(application_or_employee, JobApplication):
        app = application_or_employee
        employee = app.employee
    else:
        raise ValueError("Invalid application or employee record provided.")

    if not app:
        raise ValueError("Candidate application record is missing.")

    job = app.job
    if not job:
        raise ValueError(f"Application {app.formatted_code} is missing associated Job Posting.")

    # Determine job-specific category
    cat_key = determine_job_category(job)
    if not cat_key:
        raise OfferLetterTemplateNotFoundError(
            "No job-specific offer letter template is available for this internship. Please upload the appropriate template before generating the offer letter."
        )

    # Idempotency check: If document already exists and file exists, return it
    existing_doc = app.offer_letter_doc
    if existing_doc and existing_doc.file_path and os.path.exists(existing_doc.file_path) and not force_regenerate:
        return existing_doc, existing_doc.file_path

    # Retrieve active master template for this specific category
    active_template = get_active_offer_letter_template(cat_key)

    custom_params = custom_params or {}
    now_utc = datetime.now(timezone.utc)
    current_date_str = now_utc.strftime("%d/%m/%Y")
    
    # Calculate 7-day acceptance deadline
    default_deadline_dt = now_utc + timedelta(days=7)
    default_deadline_str = default_deadline_dt.strftime("%d %B %Y")

    # Format data from DB models
    candidate_name = app.full_name or (employee.candidate_name if employee else "Candidate")
    reference_number = app.formatted_code
    cat_info = OFFER_LETTER_CATEGORIES[cat_key]
    job_title = custom_params.get('job_title') or cat_info['default_title'] or job.title or "Intern"
    department = job.department or "Engineering"
    internship_duration = app.duration_display or (f"{job.duration.replace('_', ' ').title()}" if job.duration else "1 Month")
    emp_id_str = employee.employee_id if employee else app.formatted_code

    # Joining date
    joining_date = custom_params.get('joining_date') or custom_params.get('start_date') or "Immediate / As mutually agreed"

    # Work mode
    work_mode = custom_params.get('work_mode', job.location if job.location else "Remote")

    # Conditions
    conditions = custom_params.get(
        'conditions',
        "satisfactory verification of academic credentials and submission of government identity documentation"
    )

    # Acceptance deadline
    acceptance_deadline = custom_params.get('acceptance_deadline', default_deadline_str)

    # Responsibilities description
    responsibilities = custom_params.get(
        'responsibilities',
        job.short_description or (f"work on designated projects and deliverables for the {job_title} role at Anti-Matrix")
    )

    # Key tasks
    key_tasks = custom_params.get(
        'key_tasks',
        job.skills or (f"core technical assignments, engineering benchmarks, and team collaboration")
    )

    # Comprehensive Placeholder Dictionary (supporting both [Placeholder] and {{placeholder}} formats)
    mapping = {
        '[DD/MM/YYYY]': current_date_str,
        '{{offer_date}}': current_date_str,
        '[Candidate Name]': candidate_name,
        '{{candidate_name}}': candidate_name,
        '{{employee_name}}': candidate_name,
        '[Reference Number]': reference_number,
        '{{reference_number}}': reference_number,
        '{{application_id}}': reference_number,
        '[Job Title]': job_title,
        '{{job_title}}': job_title,
        '[1 Month / 3 Months]': internship_duration,
        '{{internship_duration}}': internship_duration,
        '[Joining Date]': joining_date,
        '{{joining_date}}': joining_date,
        '{{start_date}}': joining_date,
        '[brief description of responsibilities]': responsibilities,
        '{{responsibilities}}': responsibilities,
        '[key tasks / deliverables]': key_tasks,
        '{{key_tasks}}': key_tasks,
        '[remote / hybrid / on-site]': work_mode,
        '{{work_mode}}': work_mode,
        '[background verification / document submission / any other condition]': conditions,
        '{{conditions}}': conditions,
        '[Acceptance Deadline]': acceptance_deadline,
        '{{acceptance_deadline}}': acceptance_deadline,
        '{{employee_id}}': emp_id_str,
        '{{department}}': department
    }

    # Open active template file (CLONE - NEVER MODIFY MASTER)
    doc = docx.Document(active_template.file_path)

    # Process all standard paragraphs
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph(p, mapping)

    # Process all tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, mapping)

    # Process headers and footers across sections
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                replace_placeholders_in_paragraph(p, mapping)
        if section.footer:
            for p in section.footer.paragraphs:
                replace_placeholders_in_paragraph(p, mapping)

    # Save generated document copy
    gen_dir = os.path.join(current_app.root_path, 'uploads', 'generated_documents')
    os.makedirs(gen_dir, exist_ok=True)

    output_filename = f"{app.formatted_code}_Offer_Letter.docx"
    output_filepath = os.path.join(gen_dir, output_filename)

    doc.save(output_filepath)

    # Create or update EmployeeDocument record in DB
    emp_doc = EmployeeDocument.query.filter_by(
        application_id=app.id,
        document_type='offer_letter'
    ).first()

    if not emp_doc and employee:
        emp_doc = EmployeeDocument.query.filter_by(
            employee_id=employee.id,
            document_type='offer_letter'
        ).first()

    if not emp_doc:
        emp_doc = EmployeeDocument(
            application_id=app.id,
            employee_id=employee.id if employee else None,
            template_id=active_template.id,
            document_type='offer_letter',
            file_name=output_filename,
            file_path=output_filepath,
            status='GENERATED',
            email_status='not_sent',
            generated_at=now_utc
        )
        db.session.add(emp_doc)
    else:
        emp_doc.application_id = app.id
        if employee:
            emp_doc.employee_id = employee.id
        emp_doc.template_id = active_template.id
        emp_doc.file_name = output_filename
        emp_doc.file_path = output_filepath
        emp_doc.status = 'GENERATED'
        emp_doc.generated_at = now_utc

    db.session.commit()
    return emp_doc, output_filepath


def send_offer_letter_email(application_or_employee, start_date=None):
    """
    Sends the generated Offer Letter to the candidate's registered email with attachment.
    Enforces strict ONE-TIME send protection, Markdown-to-HTML rendering, and audit logging.
    """
    from services.email_service import send_offer_letter_shortlisted_email
    return send_offer_letter_shortlisted_email(application_or_employee, start_date=start_date)

